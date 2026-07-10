import os
import sys
import urllib.request
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from docx import Document
except ImportError:
    print("python-docx is not installed.")
    sys.exit(1)

def download_diagram():
    plantuml_code = """@startuml
left to right direction
skinparam packageStyle rectangle
skinparam usecase {
    BackgroundColor LightBlue
    BorderColor DarkBlue
}

actor "User (Pelanggan)" as u
actor Admin as a

rectangle "Sistem Barbershop Dipolka" {
  usecase "Registrasi & Login" as UC1
  usecase "Lihat Layanan & Kapster" as UC2
  usecase "Melakukan Booking" as UC3
  usecase "Melakukan Pembayaran" as UC4
  usecase "Melihat Riwayat" as UC5
  usecase "Batalkan Booking" as UC6

  usecase "Login Admin" as UC7
  usecase "Lihat Dashboard" as UC8
  usecase "Kelola Status Antrian" as UC9
  usecase "Kelola Layanan" as UC10
  usecase "Kelola Kapster" as UC11
  usecase "Lihat Laporan" as UC12
}

u --> UC1
u --> UC2
u --> UC3
u --> UC4
u --> UC5
u --> UC6

a --> UC7
a --> UC8
a --> UC9
a --> UC10
a --> UC11
a --> UC12
@enduml
"""
    url = 'https://kroki.io/plantuml/png'
    data = plantuml_code.encode('utf-8')
    headers = {
        'Content-Type': 'text/plain',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    req = urllib.request.Request(url, data=data, headers=headers)
    
    img_path = os.path.join(os.getcwd(), 'use_case_diagram.png')
    try:
        with urllib.request.urlopen(req) as response:
            with open(img_path, 'wb') as f:
                f.write(response.read())
        return img_path
    except Exception as e:
        print(f"Failed to generate diagram: {e}")
        return None

def create_use_case():
    doc = Document()
    
    # Title
    title = doc.add_heading('Dokumen Use Case - Barbershop Dipolka', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Dokumen ini berisi daftar Use Case beserta diagram untuk sistem informasi Barbershop Dipolka berbasis Laravel.')
    
    # Diagram
    img_path = download_diagram()
    if img_path and os.path.exists(img_path):
        doc.add_heading('Diagram Use Case', level=1)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_img = p_img.add_run()
        r_img.add_picture(img_path, width=Inches(6.0))
    
    # 1. Aktor Sistem
    doc.add_heading('1. Aktor Sistem', level=1)
    doc.add_paragraph('Terdapat dua aktor utama dalam sistem ini:')
    doc.add_paragraph('1. User (Pelanggan): Orang yang menggunakan aplikasi untuk memesan jadwal potong rambut (booking).', style='List Bullet')
    doc.add_paragraph('2. Admin: Pengelola barbershop yang mengatur antrian, layanan, kapster, dan melihat laporan.', style='List Bullet')
    
    # 2. Use Case: User
    doc.add_heading('2. Use Case: User (Pelanggan)', level=1)
    
    user_cases = [
        ("UC-U01: Registrasi & Login", "Pengguna mendaftarkan akun baru atau masuk ke sistem untuk dapat mengakses fitur booking."),
        ("UC-U02: Melihat Layanan & Kapster", "Pengguna dapat melihat daftar layanan (beserta harga/durasi) dan profil kapster yang sedang aktif."),
        ("UC-U03: Melakukan Booking (Antrian)", "Pengguna memilih layanan, kapster, tanggal, dan jam. Sistem akan memvalidasi jadwal agar tidak bentrok dengan pelanggan lain, serta mengecek apakah user sedang terkena penalti."),
        ("UC-U04: Melakukan Pembayaran", "Setelah booking, pengguna melakukan pembayaran secara online melalui integrasi Payment Gateway (Midtrans)."),
        ("UC-U05: Melihat Riwayat & Status Antrian", "Pengguna dapat memantau status antriannya apakah masih 'Menunggu', 'Dipanggil', 'Selesai', atau 'Batal'."),
        ("UC-U06: Membatalkan Booking", "Pengguna dapat membatalkan jadwal. Namun, jika pembatalan dilakukan lebih dari batas wajar (3 kali dalam 30 hari), sistem akan memberikan penalti (blokir fitur booking).")
    ]
    
    for title_txt, desc in user_cases:
        p = doc.add_paragraph()
        p.add_run(title_txt).bold = True
        p.add_run(f'\nDeskripsi: {desc}')
        
    # 3. Use Case: Admin
    doc.add_heading('3. Use Case: Admin', level=1)
    
    admin_cases = [
        ("UC-A01: Login Admin", "Admin melakukan proses otentikasi untuk masuk ke halaman Dashboard Admin."),
        ("UC-A02: Melihat Dashboard & Statistik", "Admin dapat melihat total pendapatan, total pelanggan, jumlah antrian yang menunggu, dan dipanggil berdasarkan filter tanggal (default: hari ini)."),
        ("UC-A03: Mengelola Status Antrian", "Admin bertugas mengubah status antrian pelanggan menjadi 'Dipanggil', 'Selesai', atau 'Batal'."),
        ("UC-A04: Mengelola Data Layanan", "Admin dapat menambah layanan baru, mengubah harga/durasi, atau menonaktifkan layanan agar tidak bisa dipilih user."),
        ("UC-A05: Mengelola Data Kapster", "Admin dapat menambah data kapster (tukang cukur), mengubah profil, atau menonaktifkannya."),
        ("UC-A06: Melihat Laporan Keuangan", "Admin dapat melihat data transaksi historis dari antrian yang telah berstatus 'Selesai'.")
    ]
    
    for title_txt, desc in admin_cases:
        p = doc.add_paragraph()
        p.add_run(title_txt).bold = True
        p.add_run(f'\nDeskripsi: {desc}')
        
    # Save the document as V2 to avoid Permission Denied if MS Word is open
    file_path = os.path.join(os.getcwd(), 'Use_Case_Barbershop_Dipolka_V2.docx')
    doc.save(file_path)
    print(f"File berhasil dibuat beserta gambar di: {file_path}")
    
    # cleanup image
    if img_path and os.path.exists(img_path):
        os.remove(img_path)

if __name__ == '__main__':
    create_use_case()
